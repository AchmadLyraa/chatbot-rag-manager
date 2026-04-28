import os
from dataclasses import dataclass
from typing import Optional
import pdfplumber
import docx

@dataclass
class LoadedDocument:
    content: str
    source: str
    type: str
    page: Optional[int] = None

def load_pdf(file_path: str) -> list[LoadedDocument]:
    docs = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(LoadedDocument(
                    content=text,
                    source=os.path.basename(file_path),
                    type="pdf",
                    page=i + 1
                ))
    print(f"Loaded {len(docs)} pages from {os.path.basename(file_path)}")
    return docs

def load_docx(file_path: str) -> list[LoadedDocument]:
    doc = docx.Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [LoadedDocument(
        content=full_text,
        source=os.path.basename(file_path),
        type="docx"
    )]

def load_txt(file_path: str) -> list[LoadedDocument]:
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    return [LoadedDocument(
        content=content,
        source=os.path.basename(file_path),
        type="txt"
    )]

def load_document(file_path: str) -> list[LoadedDocument]:
    ext = os.path.splitext(file_path)[1].lower()
    print(f"Loading: {os.path.basename(file_path)}")
    try:
        if ext == ".pdf":
            return load_pdf(file_path)
        elif ext == ".docx":
            return load_docx(file_path)
        elif ext == ".txt":
            return load_txt(file_path)
        else:
            print(f"Skipped unsupported format: {ext}")
            return []
    except Exception as e:
        print(f"Error loading {file_path}: {e}")
        return []

def load_all_documents(docs_folder: str) -> list[LoadedDocument]:
    if not os.path.exists(docs_folder):
        raise FileNotFoundError(f"Folder {docs_folder} tidak ditemukan!")

    supported = {".pdf", ".docx", ".txt"}
    files = [f for f in os.listdir(docs_folder)
             if os.path.splitext(f)[1].lower() in supported]

    print(f"\nFound {len(files)} documents in {docs_folder}\n")

    all_docs = []
    for file in files:
        file_path = os.path.join(docs_folder, file)
        all_docs.extend(load_document(file_path))

    return all_docs
